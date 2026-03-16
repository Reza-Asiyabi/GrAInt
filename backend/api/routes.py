import json
import time
from io import BytesIO
from typing import Optional

from docx import Document
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from core import database as db
from core.generator import build_client, generate_section, generate_review, revise_section
from core.prompts import SECTION_ORDER

router = APIRouter(prefix="/api")


# ─── Pydantic schemas ──────────────────────────────────────────────────────────

class ProposalInputs(BaseModel):
    topic: str
    objectives: str
    methods: str
    impact: str
    call_information: str
    references: Optional[str] = ""
    constraints: Optional[str] = ""
    timeline: Optional[str] = ""
    budget_range: Optional[str] = ""


class SectionUpdate(BaseModel):
    content: str


class RevisionRequest(BaseModel):
    feedback: str


# ─── Proposal CRUD ─────────────────────────────────────────────────────────────

@router.get("/proposals")
def list_proposals():
    return db.list_proposals()


@router.get("/proposals/{proposal_id}")
def get_proposal(proposal_id: int):
    proposal = db.get_proposal(proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    return proposal


@router.delete("/proposals/{proposal_id}", status_code=204)
def delete_proposal(proposal_id: int):
    if not db.get_proposal(proposal_id):
        raise HTTPException(status_code=404, detail="Proposal not found")
    db.delete_proposal(proposal_id)


# ─── Generation (SSE stream) ───────────────────────────────────────────────────

@router.post("/proposals/generate")
async def generate_proposal(inputs: ProposalInputs):
    """
    Server-Sent Events stream. Events emitted in order:
      { event: "start",   proposal_id: int }
      { event: "section", section: str, content: str }   ×6
      { event: "error",   section: str, message: str }   (if a section fails)
      { event: "done",    proposal_id: int }
    """
    inputs_dict = inputs.model_dump()
    proposal_id = db.create_proposal(inputs_dict)
    client = build_client()

    async def stream():
        yield f"data: {json.dumps({'event': 'start', 'proposal_id': proposal_id})}\n\n"
        db.set_status(proposal_id, "generating")

        context = {}
        for section in SECTION_ORDER:
            try:
                content = generate_section(client, section, inputs_dict, context)
                db.update_section(proposal_id, section, content)
                context[section] = content
                yield f"data: {json.dumps({'event': 'section', 'section': section, 'content': content})}\n\n"
            except Exception as exc:
                yield f"data: {json.dumps({'event': 'error', 'section': section, 'message': str(exc)})}\n\n"
            time.sleep(0.3)

        db.set_status(proposal_id, "complete")
        yield f"data: {json.dumps({'event': 'done', 'proposal_id': proposal_id})}\n\n"

    return StreamingResponse(stream(), media_type="text/event-stream")


# ─── Section editing ───────────────────────────────────────────────────────────

@router.patch("/proposals/{proposal_id}/sections/{section}")
def update_section(proposal_id: int, section: str, body: SectionUpdate):
    if not db.get_proposal(proposal_id):
        raise HTTPException(status_code=404, detail="Proposal not found")
    db.update_section(proposal_id, section, body.content)
    return {"ok": True}


@router.post("/proposals/{proposal_id}/sections/{section}/revise")
def revise(proposal_id: int, section: str, body: RevisionRequest):
    proposal = db.get_proposal(proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    client = build_client()
    current = proposal["sections"].get(section, "")
    new_content = revise_section(client, section, current, body.feedback, proposal["inputs"])
    db.update_section(proposal_id, section, new_content)
    return {"content": new_content}


# ─── Review ────────────────────────────────────────────────────────────────────

@router.post("/proposals/{proposal_id}/review")
def review_proposal(proposal_id: int):
    proposal = db.get_proposal(proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    client = build_client()
    review_text = generate_review(client, proposal["inputs"], proposal["sections"])
    db.update_review(proposal_id, review_text)
    return {"review": review_text}


# ─── Export ────────────────────────────────────────────────────────────────────

@router.get("/proposals/{proposal_id}/export")
def export_proposal(proposal_id: int, format: str = "docx"):
    proposal = db.get_proposal(proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    sections = proposal["sections"]
    title = proposal.get("title") or "Grant Proposal"

    if format == "txt":
        lines = [f"GRANT PROPOSAL\n{'=' * 60}\n"]
        for name, content in sections.items():
            lines.append(f"\n{name.replace('_', ' ').upper()}\n{'-' * 40}\n{content}\n")
        text = "\n".join(lines)
        return StreamingResponse(
            iter([text]),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="proposal_{proposal_id}.txt"'},
        )

    # DOCX
    doc = Document()
    doc.add_heading(title, 0)
    for name, content in sections.items():
        doc.add_heading(name.replace("_", " ").title(), 1)
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="proposal_{proposal_id}.docx"'},
    )
