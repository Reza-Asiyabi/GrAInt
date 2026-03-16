from openai import OpenAI
import os, json
from dotenv import load_dotenv
import json
from docx import Document


# -------------------------------
# Load API key and initialize the model
# -------------------------------
# Load API key
load_dotenv() # An OpenAI API key should be provided using a .env file

# Check if API key is loaded
if not os.getenv("OPENAI_API_KEY"):
    print("Warning: OPENAI_API_KEY not found in environment variables")

# Initialize the model with OpenAI API key
try:
    client = OpenAI(
        api_key=os.getenv("OPENAI_API_KEY"),
        base_url=os.getenv("OPENAI_BASE_URL")
    )
    print("OpenAI client initialized successfully")
except Exception as e:
    print(f"Error initializing OpenAI client: {e}")
    client = None

# -------------------------------
# Load prompt templates
# -------------------------------
with open("Prompts/Prompts.json", "r") as f:
    prompts = json.load(f)

# Define required fields for all proposal sections
required_fields = ["topic", "objectives", "methods", "impact", "call_information"]

# Collect user inputs interactively
user_inputs = {}
print("👋 Welcome to GrAInt! Let’s build your proposal step by step.\n")
for field in required_fields:
    user_inputs[field] = input(f"Please enter {field}: ")

print("\n✅ Thanks! Generating proposal sections...\n")

# -------------------------------
# Generate each section
# -------------------------------
sections_output = {}
for section, details in prompts['sections'].items():
    print(f"--- {section.upper()} ---")

    # Format the template with user inputs
    prompt_template = details["prompt"]
    final_prompt = prompt_template.format(**user_inputs)

    # Call OpenAI API
    response = client.chat.completions.create(
        model= "gpt-4o-mini", #"gpt-5-nano", #  # fast + cost effective
        messages=[{"role": "system", "content": prompts['general_writer']["prompt"].format(**user_inputs)}, {"role": "user", "content": final_prompt}],
        temperature=0.7
    )

    result = response.choices[0].message.content
    print(result + "\n")

    sections_output[section] = result


# -------------------------------
# Review, check consistency and provide feedback
# -------------------------------
# Ask user
review_choice = input("Do you want me to review the proposal, check its consistency and provide feedbacks to improve? (Y/N) ")

if review_choice == "Y":

    print('Reviewing the proposal and checking for consistency')

    # Call OpenAI API
    response = client.chat.completions.create(
        model="gpt-4o-mini",  # "gpt-5-nano", #  # fast + cost effective
        messages=[{"role": "system", "content": prompts['consistency_checker']["prompt"].format(**user_inputs)},
                  {"role": "user", "content": f"Review this document for consistency based on the system instructions and provide feedbacks and suggestions to improve it. the document: {sections_output}"}],
        temperature=0.7
    )

    result = response.choices[0].message.content
    print(result + "\n")

    sections_output['review'] = result
else:
    print("Skipped Review and Consistency Check.")

# -------------------------------
# Save results to file
# -------------------------------
def save_to_txt(filename, sections):
    with open(filename, "w", encoding="utf-8") as f:
        for section, text in sections.items():
            f.write(f"=== {section.upper()} ===\n{text}\n\n")

def save_to_docx(filename, sections):
    doc = Document()
    doc.add_heading("Research Proposal Draft", level=0)

    for section, text in sections.items():
        doc.add_heading(section.upper(), level=1)
        doc.add_paragraph(text)

    doc.save(filename)


# Ask user where to save
save_choice = input("💾 Do you want to save output as (1) TXT, (2) DOCX, or (3) Both? ")

if save_choice == "1":
    save_to_txt("Example_Outputs/proposal_draft.txt", sections_output)
    print("Saved as proposal_draft.txt ✅")
elif save_choice == "2":
    save_to_docx("Example_Outputs/proposal_draft.docx", sections_output)
    print("Saved as proposal_draft.docx ✅")
elif save_choice == "3":
    save_to_txt("Example_Outputs/proposal_draft.txt", sections_output)
    save_to_docx("Example_Outputs/proposal_draft.docx", sections_output)
    print("Saved as proposal_draft.txt and proposal_draft.docx ✅")
else:
    print("Skipped saving.")


"""
topic: developing domain-aware and interpretable AI for above ground biomass estimation from earth observation
Objectives: embed the ecological knowledge into deep learning model's architecture using intermediate interpretable feature maps to estimate less biased and more accurate AGB suing SAR and multispectral EO data.
methods: advanced deep learning architectures, satellite data, interpretable and physics-aware deep learning, using both field measured plots and spaceborne LiDAR such as GEDI as the label data
impact: the developed model will help us obtain less biased AGB maps, it will enable us to use diverse sources of training data, increase the interpretability and domain-awareness of the AI systems, hence improving the domain expert's trust in AI models. the domain-awareness also will alleviate the spurious learning.
call information: 
"""

# Please enter topic: >? AI for sustainable forest management
# Please enter objectives: >? predict tree biomass, improve monitoring, enable better conservation polic
# Please enter methods: >? deep learning, satellite data, physics-aware modeling
# Please enter impact: >? climate change mitigation, biodiversity protection

"""
Improvement Notes:

- the user should be able to provide some instructions, such as the funding call to assist the model.
- the user must be able to provide some references to the model.
- add the capability to upload an existing proposal and ask for review and then revision based on the reviews and also feedbacks from the user.
  

"""