from openai import OpenAI
import os, json
import streamlit as st
from dotenv import load_dotenv
from docx import Document
import tiktoken
from pathlib import Path
import time
from datetime import datetime
import logging
from typing import Dict, List, Optional
import re

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class TokenManager:
    """Manages token counting and optimization"""

    def __init__(self, model_name: str = "gpt-4o-mini"):
        self.model_name = model_name
        try:
            self.encoding = tiktoken.encoding_for_model(model_name)
        except KeyError:
            self.encoding = tiktoken.get_encoding("cl100k_base")

    def count_tokens(self, text: str) -> int:
        return len(self.encoding.encode(text))

    def optimize_prompt(self, prompt: str, max_tokens: int = 3000) -> str:
        """Truncate prompt if it exceeds token limit while preserving key information"""
        if self.count_tokens(prompt) <= max_tokens:
            return prompt

        # Simple truncation strategy - could be more sophisticated
        words = prompt.split()
        while self.count_tokens(' '.join(words)) > max_tokens and len(words) > 10:
            words = words[:int(len(words) * 0.9)]

        return ' '.join(words) + "..."


class ProposalGenerator:
    def __init__(self):
        load_dotenv()
        self.client = self._initialize_client()
        self.prompts = self._load_prompts()
        self.token_manager = TokenManager()
        self.sections_output = {}

    def _initialize_client(self) -> Optional[OpenAI]:
        """Initialize OpenAI client with error handling"""
        if not os.getenv("OPENAI_API_KEY"):
            logger.error("OPENAI_API_KEY not found in environment variables")
            return None

        try:
            client = OpenAI(
                api_key=os.getenv("OPENAI_API_KEY"),
                base_url=os.getenv("OPENAI_BASE_URL")
            )
            logger.info("OpenAI client initialized successfully")
            return client
        except Exception as e:
            logger.error(f"Error initializing OpenAI client: {e}")
            return None

    def _load_prompts(self) -> Dict:
        """Load prompt templates with error handling"""
        try:
            with open("Prompts/Prompts.json", "r", encoding="utf-8") as f:
                return json.load(f)
        except FileNotFoundError:
            logger.error("Prompts.json file not found")
            return {}
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing Prompts.json: {e}")
            return {}

    def collect_user_inputs(self, mode: str = "interactive") -> Dict[str, str]:
        """Collect user inputs with multiple input modes"""
        required_fields = ["topic", "objectives", "methods", "impact", "call_information"]
        optional_fields = ["references", "constraints", "timeline", "budget_range"]

        user_inputs = {}

        if mode == "interactive":
            print("🎯 Welcome to Enhanced GrAInt! Let's build your proposal step by step.\n")

            # Required fields
            for field in required_fields:
                user_inputs[field] = input(f"Please enter {field.replace('_', ' ')}: ")

            # Optional fields
            print("\n📋 Optional information (press Enter to skip):")
            for field in optional_fields:
                value = input(f"{field.replace('_', ' ')}: ")
                if value.strip():
                    user_inputs[field] = value

        elif mode == "batch":
            # Load from JSON file
            config_path = input("Enter path to configuration JSON file: ")
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    user_inputs = json.load(f)
            except Exception as e:
                logger.error(f"Error loading batch configuration: {e}")
                return {}

        return user_inputs

    def generate_section(self, section_name: str, user_inputs: Dict[str, str],
                         context: Dict[str, str] = None) -> str:
        """Generate a single section with context awareness"""
        if not self.client or section_name not in self.prompts['sections']:
            return ""

        section_details = self.prompts['sections'][section_name]
        prompt_template = section_details["prompt"]

        # Add context from previously generated sections
        if context:
            context_text = "\n\nPREVIOUS SECTIONS CONTEXT:\n"
            for ctx_section, ctx_content in context.items():
                context_text += f"{ctx_section.upper()}: {ctx_content[:200]}...\n"
            prompt_template += context_text

        # Format the template with user inputs
        final_prompt = prompt_template.format(**user_inputs)

        # Optimize for token usage
        final_prompt = self.token_manager.optimize_prompt(final_prompt)

        system_prompt = self.prompts['general_writer']["prompt"].format(**user_inputs)

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": final_prompt}
                ],
                temperature=0.7,
                max_tokens=1500  # Control output length
            )

            result = response.choices[0].message.content

            # Log token usage
            logger.info(f"Section {section_name}: ~{self.token_manager.count_tokens(final_prompt)} input tokens")

            return result

        except Exception as e:
            logger.error(f"Error generating section {section_name}: {e}")
            return f"Error generating {section_name} section."

    def generate_proposal(self, user_inputs: Dict[str, str]) -> Dict[str, str]:
        """Generate complete proposal with context awareness"""
        sections_order = ["title", "abstract", "background", "objectives", "methodology", "expected_outcomes"]
        generated_context = {}

        print("\n✨ Generating proposal sections...\n")

        for section in sections_order:
            if section in self.prompts['sections']:
                print(f"--- Generating {section.upper()} ---")

                # Pass context of previously generated sections
                result = self.generate_section(section, user_inputs, generated_context)
                print(f"{result[:200]}...\n" if len(result) > 200 else f"{result}\n")

                self.sections_output[section] = result
                generated_context[section] = result

                # Add small delay to avoid rate limiting
                time.sleep(1)

        return self.sections_output

    def review_proposal(self, sections: Dict[str, str], user_inputs: Dict[str, str]) -> str:
        """Enhanced review with specific feedback categories"""
        if not self.client:
            return "Review service unavailable."

        print('🔍 Reviewing the proposal for consistency and quality...')

        # Create structured review prompt
        review_sections = {
            "coherence": "Check for logical flow and narrative consistency",
            "technical_rigor": "Evaluate methodological soundness and feasibility",
            "impact_clarity": "Assess clarity and convincingness of expected outcomes",
            "alignment": "Check alignment with funding call requirements",
            "language_quality": "Review writing quality and academic tone"
        }

        full_document = "\n\n".join([f"=== {k.upper()} ===\n{v}" for k, v in sections.items()])

        review_prompt = f"""
        Review this research proposal comprehensively across these dimensions:
        {json.dumps(review_sections, indent=2)}

        Provide specific, actionable feedback for each dimension.
        Rate each dimension on a scale of 1-5 and provide overall recommendations.

        PROPOSAL DOCUMENT:
        {full_document}
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": self.prompts['consistency_checker']["prompt"].format(**user_inputs)},
                    {"role": "user", "content": review_prompt}
                ],
                temperature=0.3  # Lower temperature for more consistent reviews
            )

            return response.choices[0].message.content

        except Exception as e:
            logger.error(f"Error during proposal review: {e}")
            return "Error occurred during proposal review."

    def revise_section(self, section_name: str, original_content: str,
                       feedback: str, user_inputs: Dict[str, str]) -> str:
        """Revise a specific section based on feedback"""
        if not self.client:
            return original_content

        revision_prompt = f"""
        Revise the following {section_name} section based on the specific feedback provided.
        Maintain the academic tone and ensure consistency with the overall proposal.

        ORIGINAL SECTION:
        {original_content}

        FEEDBACK TO ADDRESS:
        {feedback}

        Provide only the revised section content.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": self.prompts['general_writer']["prompt"].format(**user_inputs)},
                    {"role": "user", "content": revision_prompt}
                ],
                temperature=0.5
            )

            return response.choices[0].message.content

        except Exception as e:
            logger.error(f"Error revising section {section_name}: {e}")
            return original_content

    def save_outputs(self, sections: Dict[str, str], format_choice: str = "both"):
        """Enhanced saving with multiple formats and metadata"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"proposal_draft_{timestamp}"

        # Ensure output directory exists
        output_dir = Path("Example_Outputs")
        output_dir.mkdir(exist_ok=True)

        if format_choice in ["1", "txt", "both"]:
            self._save_to_txt(output_dir / f"{base_filename}.txt", sections)
            print(f"✅ Saved as {base_filename}.txt")

        if format_choice in ["2", "docx", "both"]:
            self._save_to_docx(output_dir / f"{base_filename}.docx", sections)
            print(f"✅ Saved as {base_filename}.docx")

        if format_choice in ["3", "json", "both"]:
            self._save_to_json(output_dir / f"{base_filename}.json", sections)
            print(f"✅ Saved as {base_filename}.json")

    def _save_to_txt(self, filename: Path, sections: Dict[str, str]):
        """Save to text with metadata"""
        with open(filename, "w", encoding="utf-8") as f:
            f.write(f"Generated by GrAInt on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 60 + "\n\n")

            for section, text in sections.items():
                f.write(f"=== {section.upper()} ===\n{text}\n\n")

    def _save_to_docx(self, filename: Path, sections: Dict[str, str]):
        """Enhanced DOCX saving with better formatting"""
        doc = Document()

        # Add title
        doc.add_heading("Research Proposal Draft", level=0)

        # Add metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f"Generated by GrAInt on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

        for section, text in sections.items():
            doc.add_heading(section.replace('_', ' ').title(), level=1)

            # Handle potential formatting in the text
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

        doc.save(filename)

    def _save_to_json(self, filename: Path, sections: Dict[str, str]):
        """Save as structured JSON with metadata"""
        output_data = {
            "metadata": {
                "generated_by": "GrAInt",
                "timestamp": datetime.now().isoformat(),
                "version": "2.0"
            },
            "sections": sections
        }

        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)


def main():
    """Main application flow with enhanced user experience"""
    generator = ProposalGenerator()

    if not generator.client:
        print("❌ Cannot proceed without OpenAI client. Please check your API key.")
        return

    # Input mode selection
    print("🚀 Enhanced GrAInt - Grant Proposal Assistant")
    print("Choose input mode:")
    print("1. Interactive (step-by-step)")
    print("2. Batch (from JSON file)")

    mode_choice = input("Enter choice (1/2): ").strip()
    input_mode = "interactive" if mode_choice == "1" else "batch"

    # Collect inputs
    user_inputs = generator.collect_user_inputs(input_mode)
    if not user_inputs:
        print("❌ Failed to collect user inputs.")
        return

    # Generate proposal
    sections = generator.generate_proposal(user_inputs)

    # Review option
    review_choice = input("\n🔍 Do you want me to review the proposal? (Y/N): ").upper()
    if review_choice == "Y":
        review_feedback = generator.review_proposal(sections, user_inputs)
        print("\n📋 REVIEW FEEDBACK:")
        print(review_feedback)
        sections['review'] = review_feedback

        # Revision option
        revise_choice = input("\n✏️ Do you want to revise any sections based on feedback? (Y/N): ").upper()
        if revise_choice == "Y":
            print("Available sections:", list(sections.keys()))
            section_to_revise = input("Which section would you like to revise? ")

            if section_to_revise in sections:
                revised_content = generator.revise_section(
                    section_to_revise,
                    sections[section_to_revise],
                    review_feedback,
                    user_inputs
                )
                sections[section_to_revise] = revised_content
                print(f"✅ Revised {section_to_revise} section")

    # Save options
    print("\n💾 Save options:")
    print("1. TXT format")
    print("2. DOCX format")
    print("3. JSON format")
    print("4. All formats")

    save_choice = input("Enter choice (1-4): ").strip()
    format_map = {"1": "txt", "2": "docx", "3": "json", "4": "both"}

    generator.save_outputs(sections, format_map.get(save_choice, "both"))

    print("\n🎉 Proposal generation completed!")


if __name__ == "__main__":
    main()