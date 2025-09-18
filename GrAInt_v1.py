import streamlit as st
import json
import os
from datetime import datetime
from pathlib import Path
import time
import io
from typing import Dict, List, Optional
import pandas as pd

from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
import tiktoken
import logging

# Configure page
st.set_page_config(
    page_title="GrAInt - AI Grant Proposal Assistant",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        color: #2c5aa0;
        border-bottom: 2px solid #1f77b4;
        padding-bottom: 0.5rem;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
    .error-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        margin: 1rem 0;
    }
    .info-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #cce7ff;
        border: 1px solid #99d6ff;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'proposal_sections' not in st.session_state:
    st.session_state.proposal_sections = {}
if 'user_inputs' not in st.session_state:
    st.session_state.user_inputs = {}
if 'generation_complete' not in st.session_state:
    st.session_state.generation_complete = False
if 'review_feedback' not in st.session_state:
    st.session_state.review_feedback = ""


class StreamlitProposalGenerator:
    def __init__(self):
        load_dotenv()
        self.client = self._initialize_client()
        self.prompts = self._load_prompts()
        self.token_manager = TokenManager() if 'TokenManager' in globals() else None

    def _initialize_client(self) -> Optional[OpenAI]:
        """Initialize OpenAI client with error handling"""
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            st.error("⚠️ OpenAI API key not found. Please set OPENAI_API_KEY in your environment variables.")
            return None

        try:
            client = OpenAI(
                api_key=api_key,
                base_url=os.getenv("OPENAI_BASE_URL")
            )
            return client
        except Exception as e:
            st.error(f"❌ Error initializing OpenAI client: {e}")
            return None

    def _load_prompts(self) -> Dict:
        """Load prompt templates with error handling"""
        try:
            # Try to load from the prompts file
            if os.path.exists("Prompts/Prompts.json"):
                with open("Prompts/Prompts.json", "r", encoding="utf-8") as f:
                    return json.load(f)
            else:
                # Fallback to embedded prompts
                return self._get_default_prompts()
        except Exception as e:
            st.warning(f"⚠️ Could not load prompts file: {e}. Using default prompts.")
            return self._get_default_prompts()

    def _get_default_prompts(self) -> Dict:
        """Default prompts if file is not available"""
        return {
            "general_writer": {
                "prompt": "You are an expert academic writer with over 20 years of experience crafting successful UKRI and Horizon Europe proposals. You write in a clear, persuasive, and human academic style that avoids artificial or generic language. Your writing is ambitious yet realistic, demonstrates originality and impact, and consistently aligns with funding body priorities. Every section should connect smoothly, maintaining a coherent and compelling narrative across the proposal. Extra information about the call: {call_information}"
            },
            "consistency_checker": {
                "prompt": "You are a senior research evaluator reviewing a full draft of a research proposal. You provide constructive, specific, and actionable feedback to strengthen proposals and maximize their chances of success. Check the entire document for consistency in tone, terminology, and narrative flow. Ensure that the objectives, methodology, outcomes, and impact are logically aligned. Extra information about the call: {call_information}"
            },
            "sections": {
                "title": {
                    "prompt": "Propose a title for a research proposal on: {topic}. Objectives: {objectives}. Methodology: {methods}. Impact: {impact}. The title must be under 15 words, formal yet engaging, and immediately appealing to reviewers."
                },
                "abstract": {
                    "prompt": "Write a 250-word abstract for a research proposal. Topic: {topic}. Objectives: {objectives}. Methodology: {methods}. Impact: {impact}. Emphasize novelty, expected outcomes, and impact."
                },
                "background": {
                    "prompt": "Draft the background section for: Topic: {topic}. Objectives: {objectives}. Methodology: {methods}. Impact: {impact}. Explain why this research matters, highlight current state, and identify gaps."
                },
                "objectives": {
                    "prompt": "Write 3-5 clear, measurable objectives based on: Objectives: {objectives}. Methodology: {methods}. Impact: {impact}. Phrase as concrete research goals."
                },
                "methodology": {
                    "prompt": "Write the methodology section. Describe data, methods, tools, and innovative aspects. Objectives: {objectives}. Methodology: {methods}. Highlight novelty and feasibility."
                },
                "expected_outcomes": {
                    "prompt": "Write expected outcomes section. Objectives: {objectives}. Impact: {impact}. Describe anticipated contributions, societal relevance, and alignment with funding priorities."
                }
            }
        }

    def generate_section(self, section_name: str, user_inputs: Dict[str, str],
                         context: Dict[str, str] = None) -> str:
        """Generate a single section with progress tracking"""
        if not self.client or section_name not in self.prompts['sections']:
            return "Error: Unable to generate section"

        section_details = self.prompts['sections'][section_name]
        prompt_template = section_details["prompt"]

        # Add context from previously generated sections
        if context:
            context_text = "\n\nPREVIOUS SECTIONS CONTEXT:\n"
            for ctx_section, ctx_content in context.items():
                context_text += f"{ctx_section.upper()}: {ctx_content[:200]}...\n"
            prompt_template += context_text

        try:
            final_prompt = prompt_template.format(**user_inputs)
            system_prompt = self.prompts['general_writer']["prompt"].format(**user_inputs)

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": final_prompt}
                ],
                temperature=0.7,
                max_tokens=1500
            )

            return response.choices[0].message.content

        except Exception as e:
            return f"Error generating {section_name}: {str(e)}"

    def review_proposal(self, sections: Dict[str, str], user_inputs: Dict[str, str]) -> str:
        """Review the complete proposal"""
        if not self.client:
            return "Review service unavailable."

        full_document = "\n\n".join([f"=== {k.upper()} ===\n{v}" for k, v in sections.items() if k != 'review'])

        review_prompt = f"""
        Review this research proposal across these dimensions:
        1. Coherence: Logical flow and narrative consistency
        2. Technical rigor: Methodological soundness and feasibility  
        3. Impact clarity: Convincingness of expected outcomes
        4. Alignment: Fit with funding call requirements
        5. Language quality: Writing quality and academic tone

        Provide specific, actionable feedback for each dimension.

        PROPOSAL DOCUMENT:
        {full_document}
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": self.prompts['consistency_checker']["prompt"].format(**user_inputs)},
                    {"role": "user", "content": review_prompt}
                ],
                temperature=0.3
            )

            return response.choices[0].message.content

        except Exception as e:
            return f"Error during review: {str(e)}"


class TokenManager:
    """Simple token manager for the demo"""

    def __init__(self):
        try:
            self.encoding = tiktoken.get_encoding("cl100k_base")
        except:
            self.encoding = None

    def count_tokens(self, text: str) -> int:
        if self.encoding:
            return len(self.encoding.encode(text))
        return len(text.split()) * 1.3  # Rough estimate


def main():
    """Main Streamlit application"""

    # Header
    st.markdown('<h1 class="main-header">🎯 GrAInt</h1>', unsafe_allow_html=True)
    st.markdown(
        '<p style="text-align: center; font-size: 1.2rem; color: #666;">AI-Powered Grant Proposal Assistant</p>',
        unsafe_allow_html=True)

    # Initialize the generator
    if 'generator' not in st.session_state:
        st.session_state.generator = StreamlitProposalGenerator()

    # Sidebar for navigation and settings
    st.sidebar.title("Navigation")

    # Check API connection
    if st.session_state.generator.client is None:
        st.sidebar.error("❌ OpenAI client not initialized")
        st.error("Please configure your OpenAI API key in the environment variables and restart the application.")
        return
    else:
        st.sidebar.success("✅ OpenAI client connected")

    # Navigation
    page = st.sidebar.radio("Choose a section:", [
        "📝 Input Information",
        "🔧 Generate Proposal",
        "🔍 Review & Revise",
        "💾 Export Results"
    ])

    if page == "📝 Input Information":
        show_input_page()
    elif page == "🔧 Generate Proposal":
        show_generation_page()
    elif page == "🔍 Review & Revise":
        show_review_page()
    elif page == "💾 Export Results":
        show_export_page()


def show_input_page():
    """Input information page"""
    st.markdown('<h2 class="section-header">📝 Proposal Information</h2>', unsafe_allow_html=True)

    # Create tabs for different input methods
    tab1, tab2, tab3 = st.tabs(["Manual Input", "Upload JSON", "Load Template"])

    with tab1:
        st.markdown("### Required Information")

        col1, col2 = st.columns(2)

        with col1:
            topic = st.text_area(
                "Research Topic",
                value=st.session_state.user_inputs.get('topic', ''),
                height=100,
                help="Describe your research topic concisely"
            )

            objectives = st.text_area(
                "Research Objectives",
                value=st.session_state.user_inputs.get('objectives', ''),
                height=150,
                help="What are you trying to achieve? List 3-5 main objectives"
            )

            methods = st.text_area(
                "Methodology",
                value=st.session_state.user_inputs.get('methods', ''),
                height=150,
                help="How will you conduct the research? Include data, tools, approaches"
            )

        with col2:
            impact = st.text_area(
                "Expected Impact",
                value=st.session_state.user_inputs.get('impact', ''),
                height=150,
                help="What will be the broader impact of your research?"
            )

            call_information = st.text_area(
                "Funding Call Information",
                value=st.session_state.user_inputs.get('call_information', ''),
                height=100,
                help="Information about the specific funding call (optional)"
            )

        st.markdown("### Optional Information")

        col3, col4 = st.columns(2)
        with col3:
            references = st.text_area(
                "Key References",
                value=st.session_state.user_inputs.get('references', ''),
                height=100,
                help="Important references or literature (optional)"
            )

            timeline = st.text_input(
                "Project Timeline",
                value=st.session_state.user_inputs.get('timeline', ''),
                help="e.g., '36 months', '3 years' (optional)"
            )

        with col4:
            budget_range = st.text_input(
                "Budget Range",
                value=st.session_state.user_inputs.get('budget_range', ''),
                help="e.g., '£500K-£1M' (optional)"
            )

            constraints = st.text_area(
                "Constraints/Requirements",
                value=st.session_state.user_inputs.get('constraints', ''),
                height=100,
                help="Any specific constraints or requirements (optional)"
            )

        # Save inputs button
        if st.button("💾 Save Information", type="primary"):
            st.session_state.user_inputs = {
                'topic': topic,
                'objectives': objectives,
                'methods': methods,
                'impact': impact,
                'call_information': call_information,
                'references': references,
                'timeline': timeline,
                'budget_range': budget_range,
                'constraints': constraints
            }
            st.success("✅ Information saved successfully!")

    with tab2:
        st.markdown("### Upload Configuration File")
        uploaded_file = st.file_uploader("Choose a JSON file", type=['json'])

        if uploaded_file is not None:
            try:
                config = json.load(uploaded_file)
                st.session_state.user_inputs = config
                st.success("✅ Configuration loaded successfully!")
                st.json(config)
            except Exception as e:
                st.error(f"❌ Error loading file: {e}")

    with tab3:
        st.markdown("### Load from Template")
        template_options = {
            "AI/ML Research": {
                'topic': "Machine learning for climate prediction",
                'objectives': "Develop novel ML algorithms, improve prediction accuracy, create interpretable models",
                'methods': "Deep learning, ensemble methods, satellite data analysis",
                'impact': "Better climate predictions, policy support, scientific advancement",
                'call_information': "UKRI AI for Science call"
            },
            "Biotechnology": {
                'topic': "CRISPR gene editing for disease treatment",
                'objectives': "Improve gene editing precision, develop therapeutic applications, ensure safety",
                'methods': "CRISPR-Cas9, cell culture, animal models, clinical trials",
                'impact': "New treatments, reduced healthcare costs, improved patient outcomes",
                'call_information': "Medical Research Council funding"
            },
            "Social Sciences": {
                'topic': "Digital divide and social inequality",
                'objectives': "Analyze digital access patterns, identify inequality factors, propose interventions",
                'methods': "Survey research, statistical analysis, policy analysis, case studies",
                'impact': "Reduced digital inequality, better policy, social cohesion",
                'call_information': "ESRC digital society call"
            }
        }

        template_choice = st.selectbox("Choose a template:", list(template_options.keys()))

        if st.button("Load Template"):
            st.session_state.user_inputs = template_options[template_choice]
            st.success(f"✅ {template_choice} template loaded!")


def show_generation_page():
    """Proposal generation page"""
    st.markdown('<h2 class="section-header">🔧 Generate Proposal</h2>', unsafe_allow_html=True)

    # Check if inputs are available
    if not st.session_state.user_inputs.get('topic'):
        st.warning("⚠️ Please provide at least the research topic in the Input Information section.")
        return

    # Show current inputs summary
    with st.expander("📋 Current Information Summary", expanded=False):
        for key, value in st.session_state.user_inputs.items():
            if value:
                st.write(f"**{key.replace('_', ' ').title()}:** {value[:100]}{'...' if len(value) > 100 else ''}")

    # Generation options
    st.markdown("### Generation Settings")
    col1, col2, col3 = st.columns(3)

    with col1:
        sections_to_generate = st.multiselect(
            "Sections to Generate:",
            ["title", "abstract", "background", "objectives", "methodology", "expected_outcomes"],
            default=["title", "abstract", "background", "objectives", "methodology", "expected_outcomes"]
        )

    with col2:
        temperature = st.slider("Creativity Level", 0.1, 1.0, 0.7, 0.1,
                                help="Higher values = more creative, Lower values = more conservative")

    with col3:
        context_aware = st.checkbox("Context-Aware Generation", value=True,
                                    help="Each section considers previously generated sections")

    # Generation button
    if st.button("🚀 Generate Proposal", type="primary"):
        generate_proposal_sections(sections_to_generate, temperature, context_aware)

    # Display generated sections
    if st.session_state.proposal_sections:
        st.markdown("### Generated Sections")

        for section in sections_to_generate:
            if section in st.session_state.proposal_sections:
                with st.expander(f"📄 {section.replace('_', ' ').title()}", expanded=True):
                    st.write(st.session_state.proposal_sections[section])

                    # Edit button for each section
                    if st.button(f"✏️ Edit {section.title()}", key=f"edit_{section}"):
                        edit_section_modal(section)


def generate_proposal_sections(sections_to_generate, temperature, context_aware):
    """Generate proposal sections with progress tracking"""
    progress_bar = st.progress(0)
    status_text = st.empty()
    generated_context = {}

    for i, section in enumerate(sections_to_generate):
        status_text.text(f"Generating {section.replace('_', ' ').title()}...")

        # Generate section
        context = generated_context if context_aware else None
        result = st.session_state.generator.generate_section(section, st.session_state.user_inputs, context)

        # Store result
        st.session_state.proposal_sections[section] = result
        if context_aware:
            generated_context[section] = result

        # Update progress
        progress_bar.progress((i + 1) / len(sections_to_generate))
        time.sleep(0.5)  # Small delay for visual feedback

    status_text.text("✅ Generation complete!")
    st.session_state.generation_complete = True
    st.success("🎉 Proposal sections generated successfully!")


def edit_section_modal(section_name):
    """Show edit modal for a section"""
    st.markdown(f"### Edit {section_name.replace('_', ' ').title()}")

    current_content = st.session_state.proposal_sections.get(section_name, "")
    edited_content = st.text_area(
        "Content:",
        value=current_content,
        height=300,
        key=f"edit_area_{section_name}"
    )

    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save Changes", key=f"save_{section_name}"):
            st.session_state.proposal_sections[section_name] = edited_content
            st.success("Changes saved!")

    with col2:
        if st.button("🔄 Regenerate", key=f"regen_{section_name}"):
            with st.spinner("Regenerating..."):
                new_content = st.session_state.generator.generate_section(
                    section_name,
                    st.session_state.user_inputs
                )
                st.session_state.proposal_sections[section_name] = new_content
                st.success("Section regenerated!")
                st.experimental_rerun()


def show_review_page():
    """Review and revision page"""
    st.markdown('<h2 class="section-header">🔍 Review & Revise</h2>', unsafe_allow_html=True)

    if not st.session_state.proposal_sections:
        st.warning("⚠️ Please generate proposal sections first.")
        return

    # Review button
    if st.button("🔍 Generate Review", type="primary"):
        with st.spinner("Reviewing proposal..."):
            review = st.session_state.generator.review_proposal(
                st.session_state.proposal_sections,
                st.session_state.user_inputs
            )
            st.session_state.review_feedback = review

    # Display review
    if st.session_state.review_feedback:
        st.markdown("### Review Feedback")
        st.write(st.session_state.review_feedback)

        # Revision options
        st.markdown("### Revision Tools")

        section_to_revise = st.selectbox(
            "Select section to revise:",
            list(st.session_state.proposal_sections.keys())
        )

        specific_feedback = st.text_area(
            "Specific feedback for revision:",
            help="Enter specific points you want addressed in the revision"
        )

        if st.button("🔄 Revise Section"):
            with st.spinner("Revising section..."):
                # Here you could implement section-specific revision
                st.success("Section revised! (Feature to be implemented)")

    # Quality metrics
    st.markdown("### Quality Metrics")
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("Word Count", calculate_word_count())

    with col2:
        st.metric("Sections", len(st.session_state.proposal_sections))

    with col3:
        st.metric("Completeness", f"{calculate_completeness()}%")

    with col4:
        if st.session_state.generator.token_manager:
            tokens = sum([st.session_state.generator.token_manager.count_tokens(content)
                          for content in st.session_state.proposal_sections.values()])
            st.metric("Est. Tokens", f"{tokens:,}")


def calculate_word_count():
    """Calculate total word count of generated sections"""
    total_words = sum([len(content.split()) for content in st.session_state.proposal_sections.values()])
    return f"{total_words:,}"


def calculate_completeness():
    """Calculate proposal completeness percentage"""
    expected_sections = ["title", "abstract", "background", "objectives", "methodology", "expected_outcomes"]
    completed = len([s for s in expected_sections if s in st.session_state.proposal_sections])
    return int((completed / len(expected_sections)) * 100)


def show_export_page():
    """Export and download page"""
    st.markdown('<h2 class="section-header">💾 Export Results</h2>', unsafe_allow_html=True)

    if not st.session_state.proposal_sections:
        st.warning("⚠️ No proposal sections to export. Please generate sections first.")
        return

    # Export options
    st.markdown("### Export Options")

    col1, col2 = st.columns(2)

    with col1:
        # Text format
        st.subheader("📄 Text Format")
        text_content = generate_text_export()
        st.download_button(
            label="Download as TXT",
            data=text_content,
            file_name=f"proposal_draft_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            mime="text/plain"
        )

    with col2:
        # JSON format
        st.subheader("📊 JSON Format")
        json_content = generate_json_export()
        st.download_button(
            label="Download as JSON",
            data=json_content,
            file_name=f"proposal_draft_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json"
        )

    # Word document
    st.subheader("📝 Word Document")
    if st.button("📥 Generate Word Document"):
        doc_bytes = generate_docx_export()
        st.download_button(
            label="Download as DOCX",
            data=doc_bytes,
            file_name=f"proposal_draft_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Preview
    st.markdown("### Preview")
    with st.expander("📖 Full Proposal Preview", expanded=False):
        for section, content in st.session_state.proposal_sections.items():
            st.markdown(f"### {section.replace('_', ' ').title()}")
            st.write(content)
            st.markdown("---")


def generate_text_export():
    """Generate text format export"""
    content = f"Generated by GrAInt on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    content += "=" * 60 + "\n\n"

    for section, text in st.session_state.proposal_sections.items():
        content += f"=== {section.upper().replace('_', ' ')} ===\n{text}\n\n"

    return content


def generate_json_export():
    """Generate JSON format export"""
    output_data = {
        "metadata": {
            "generated_by": "GrAInt",
            "timestamp": datetime.now().isoformat(),
            "user_inputs": st.session_state.user_inputs
        },
        "sections": st.session_state.proposal_sections
    }

    return json.dumps(output_data, indent=2, ensure_ascii=False)


def generate_docx_export():
    """Generate DOCX format export"""
    doc = Document()

    # Add title
    doc.add_heading("Research Proposal Draft", level=0)

    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Generated by GrAInt on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    for section, text in st.session_state.proposal_sections.items():
        doc.add_heading(section.replace('_', ' ').title(), level=1)

        # Handle potential formatting in the text
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()


if __name__ == "__main__":
    main()